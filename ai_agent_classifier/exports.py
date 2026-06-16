import io
from datetime import datetime

from openpyxl import Workbook
from openpyxl.styles import (Alignment, Border, Font, PatternFill, Side)
from openpyxl.utils import get_column_letter

from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm, mm
from reportlab.lib.colors import HexColor, white, black, Color
from reportlab.platypus import (
    SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer,
    PageBreak, HRFlowable, KeepTogether,
)
from reportlab.platypus.flowables import Flowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY


# ─── Shared constants ─────────────────────────────────────────────────────────

STAGES = [
    ("idea-gen",    "Idea Generation"),
    ("idea-assess", "Idea Assessment"),
    ("decision",    "Decision Point"),
    ("execution",   "Execution"),
    ("monitoring",  "Monitoring"),
    ("compliance",  "Compliance"),
    ("stakeholder", "Stakeholder Management"),
]
STAGE_KEYS   = [s[0] for s in STAGES]
STAGE_LABELS = [s[1] for s in STAGES]

STAGE_ABBREV = {
    "idea-gen":    "Idea Gen.",
    "idea-assess": "Idea Assess.",
    "decision":    "Decision",
    "execution":   "Execution",
    "monitoring":  "Monitoring",
    "compliance":  "Compliance",
    "stakeholder": "Stakeholder",
}

COMPLEXITIES = [
    ("black",      "Black Swan (Unknown unknowns)"),
    ("dark-grey",  "Dark Grey Swan (Scenario analysis / causation-based)"),
    ("light-grey", "Light Grey Swan (Noise reduction / optimisation heuristics)"),
    ("white",      "White Swan (Statistical / deterministic)"),
]
COMPLEXITY_KEYS   = [c[0] for c in COMPLEXITIES]
COMPLEXITY_LABELS = [c[1] for c in COMPLEXITIES]

COMPLEXITY_SHORT = {
    "black":      "BLACK SWAN",
    "dark-grey":  "DARK GREY SWAN",
    "light-grey": "LIGHT GREY SWAN",
    "white":      "WHITE SWAN",
}

ADVANTAGES = [
    ("behavioral",    "Behavioral"),
    ("analytical",    "Analytical"),
    ("informational", "Informational"),
]
ADVANTAGE_KEYS   = [a[0] for a in ADVANTAGES]
ADVANTAGE_LABELS = [a[1] for a in ADVANTAGES]

COMPLEXITY_HINTS = {
    "white":      "Statistical models / known distribution",
    "light-grey": "Noise reduction / optimisation heuristics",
    "dark-grey":  "Causal identification / scarce information",
    "black":      "Deep uncertainty / robustness over optimality",
}

COLOR_BEH_SOLID = "D85A30"
COLOR_ANA_SOLID = "1D9E75"
COLOR_INF_SOLID = "EF9F27"
COLOR_BEH_TINT  = "FAE6E1"
COLOR_ANA_TINT  = "E0F3EE"
COLOR_INF_TINT  = "FDF3E0"

COMPLEXITY_COLORS = {
    "black":      "C0392B",
    "dark-grey":  "D4AC0D",
    "light-grey": "2980B9",
    "white":      "27AE60",
}

COMPLEXITY_BADGE_BG = {
    "black":      "000000",
    "dark-grey":  "555555",
    "light-grey": "AAAAAA",
    "white":      "FFFFFF",
}

THIN_SIDE   = Side(style="thin", color="BDBDBD")
THIN_BORDER = Border(left=THIN_SIDE, right=THIN_SIDE, top=THIN_SIDE, bottom=THIN_SIDE)


def _advantage_tint(adv_key):
    return {"behavioral": COLOR_BEH_TINT, "analytical": COLOR_ANA_TINT,
            "informational": COLOR_INF_TINT}.get(adv_key, "FFFFFF")


def _advantage_solid(adv_key):
    return {"behavioral": COLOR_BEH_SOLID, "analytical": COLOR_ANA_SOLID,
            "informational": COLOR_INF_SOLID}.get(adv_key, "888888")


def _sort_by_stage(agents):
    """Sort by first pipeline stage, then complexity, then name."""
    cx_ord    = {"black": 0, "dark-grey": 1, "light-grey": 2, "white": 3}
    st_order  = {sk: i for i, sk in enumerate(STAGE_KEYS)}

    def key(a):
        first = min((st_order.get(s, 99) for s in a.stages_list), default=99)
        return (first, cx_ord.get(a.complexity, 9), a.name.lower())

    return sorted(agents, key=key)


# ─── Excel helpers ────────────────────────────────────────────────────────────

def _fill(hex_color):
    return PatternFill("solid", fgColor=hex_color)


def _font(bold=False, size=10, color=None, italic=False):
    kw = {"bold": bold, "size": size, "italic": italic}
    if color:
        kw["color"] = color
    return Font(**kw)


def _align(h="left", v="top", wrap=True):
    return Alignment(horizontal=h, vertical=v, wrap_text=wrap)


# ─── Excel export ────────────────────────────────────────────────────────────

def build_excel(agents):
    wb = Workbook()
    _sheet_matrix(wb, agents)
    _sheet_details(wb, agents)
    _sheet_summary(wb, agents)
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


def _sheet_matrix(wb, agents):
    ws = wb.active
    ws.title = "AI Agents Classification Matrix"

    ws.column_dimensions["A"].width = 32
    ws.column_dimensions["B"].width = 4
    for col_idx in range(3, 10):
        ws.column_dimensions[get_column_letter(col_idx)].width = 22
    ws.column_dimensions["J"].width = 6
    ws.column_dimensions["K"].width = 8

    ws.row_dimensions[1].height = 28
    ws["B1"] = "Complexity Range"
    ws["B1"].font = _font(bold=True, size=11)
    ws.merge_cells("C1:I1")
    ws["C1"] = "Investment Process Value Chain"
    ws["C1"].font      = _font(bold=True, size=11)
    ws["C1"].alignment = _align(h="center", v="center")
    ws["K1"] = "Total"
    ws["K1"].font      = _font(bold=True, size=11)
    ws["K1"].alignment = _align(h="center")

    ws.row_dimensions[2].height = 28
    grey_fill = _fill("E8E8E8")
    for col_idx, label in enumerate(STAGE_LABELS, start=3):
        cell           = ws.cell(row=2, column=col_idx, value=label)
        cell.font      = _font(bold=True, size=10)
        cell.fill      = grey_fill
        cell.alignment = _align(h="center", v="center")
        cell.border    = THIN_BORDER
    for col_idx in [1, 2, 10, 11]:
        ws.cell(row=2, column=col_idx).fill = grey_fill

    row_colors = {"black": "FEF2F2", "dark-grey": "FEFCE8",
                  "light-grey": "EFF6FF", "white": "F0FDF4"}
    row = 3

    for cx_key, cx_label in COMPLEXITIES:
        cx_agents    = [a for a in agents if a.complexity == cx_key]
        cx_start_row = row
        bg           = row_colors.get(cx_key, "FFFFFF")

        for adv_idx, (adv_key, _) in enumerate(ADVANTAGES):
            ws.row_dimensions[row].height = 42
            tint_hex  = _advantage_tint(adv_key)
            solid_hex = _advantage_solid(adv_key)

            if adv_idx == 0:
                ws.cell(row=row, column=1, value=cx_label).font      = _font(bold=True, size=10)
                ws.cell(row=row, column=1).alignment = _align(h="left", v="center", wrap=True)

            badge_cell           = ws.cell(row=row, column=10, value=adv_key[:3].upper())
            badge_cell.fill      = _fill(solid_hex)
            badge_cell.font      = _font(bold=True, size=10, color="FFFFFF")
            badge_cell.alignment = _align(h="center", v="center", wrap=False)
            badge_cell.border    = THIN_BORDER

            for col_idx in range(1, 12):
                c = ws.cell(row=row, column=col_idx)
                if col_idx != 10:
                    c.fill = _fill(bg)
                c.border = THIN_BORDER

            for stage_idx, (stage_key, _) in enumerate(STAGES):
                col_idx  = stage_idx + 3
                matching = [a for a in cx_agents
                            if a.advantage == adv_key and stage_key in a.stages_list]
                cell = ws.cell(row=row, column=col_idx)
                if matching:
                    cell.value = f"{len(matching)}  {'  '.join(a.name for a in matching)}"
                cell.fill      = _fill(tint_hex)
                cell.alignment = _align(h="left", v="top", wrap=True)
                cell.border    = THIN_BORDER
                cell.font      = _font(size=9)

            row += 1

        if cx_start_row + 2 <= row - 1:
            ws.merge_cells(start_row=cx_start_row, start_column=1,
                           end_row=row - 1, end_column=1)

        total_cx   = len(set(a.id for a in cx_agents))
        total_cell = ws.cell(row=cx_start_row, column=11, value=total_cx)
        total_cell.font      = _font(bold=True, size=11)
        total_cell.alignment = _align(h="center", v="center")
        if cx_start_row + 2 <= row - 1:
            ws.merge_cells(start_row=cx_start_row, start_column=11,
                           end_row=row - 1, end_column=11)

    ws.row_dimensions[row].height = 8
    row += 1

    ws.row_dimensions[row].height = 28
    ws.cell(row=row, column=1, value="Total").font   = _font(bold=True, size=10)
    ws.cell(row=row, column=1).fill   = _fill("F5F5F5")
    ws.cell(row=row, column=1).border = THIN_BORDER
    grand_total = len(set(a.id for a in agents))
    for stage_idx, (stage_key, _) in enumerate(STAGES):
        col_idx = stage_idx + 3
        count   = len(set(a.id for a in agents if stage_key in a.stages_list))
        cell    = ws.cell(row=row, column=col_idx, value=count)
        cell.font      = _font(bold=True, size=10)
        cell.fill      = _fill("F5F5F5")
        cell.alignment = _align(h="center")
        cell.border    = THIN_BORDER
    ws.cell(row=row, column=11, value=grand_total).font      = _font(bold=True, size=11)
    ws.cell(row=row, column=11).fill      = _fill("F5F5F5")
    ws.cell(row=row, column=11).alignment = _align(h="center")
    ws.cell(row=row, column=11).border    = THIN_BORDER
    row += 1

    ws.row_dimensions[row].height = 28
    ws.cell(row=row, column=1, value="By advantage").font   = _font(bold=True, size=10)
    ws.cell(row=row, column=1).border = THIN_BORDER
    beh_count = len([a for a in agents if a.advantage == "behavioral"])
    ana_count = len([a for a in agents if a.advantage == "analytical"])
    inf_count = len([a for a in agents if a.advantage == "informational"])
    ws.merge_cells(start_row=row, start_column=3, end_row=row, end_column=11)
    color_cell           = ws.cell(row=row, column=3,
                                   value=f"BEH: {beh_count}          ANA: {ana_count}          INF: {inf_count}")
    color_cell.alignment = _align(h="center", v="center")
    color_cell.border    = THIN_BORDER
    ws.freeze_panes = "C3"


def _sheet_details(wb, agents):
    ws      = wb.create_sheet("Agent Details")
    headers = ["Name", "URL", "Type", "Stage(s)", "Complexity",
               "Advantage", "Autonomy", "Description", "Key Features", "Rationale"]
    hfill   = _fill("E8E8E8")
    for col_idx, h in enumerate(headers, 1):
        cell           = ws.cell(row=1, column=col_idx, value=h)
        cell.font      = _font(bold=True, size=10)
        cell.fill      = hfill
        cell.border    = THIN_BORDER
        cell.alignment = _align(h="center")

    for row_idx, agent in enumerate(_sort_by_stage(agents), 2):
        bg   = "FAFAFA" if row_idx % 2 == 0 else "FFFFFF"
        vals = [
            agent.name,
            agent.url or "",
            agent.agent_type or "",
            ", ".join(dict(STAGES).get(s, s) for s in agent.stages_list),
            dict(zip(COMPLEXITY_KEYS, COMPLEXITY_LABELS)).get(agent.complexity, ""),
            agent.advantage or "",
            agent.autonomy or "",
            agent.description or "",
            " | ".join(agent.features_list),
            agent.rationale or "",
        ]
        for col_idx, val in enumerate(vals, 1):
            cell           = ws.cell(row=row_idx, column=col_idx, value=val)
            cell.fill      = _fill(bg)
            cell.border    = THIN_BORDER
            cell.alignment = _align(h="left", v="top", wrap=True)
            cell.font      = _font(size=9)
            if col_idx == 2 and val:
                cell.value = f'=HYPERLINK("{val}", "{val}")'
                cell.font  = Font(size=9, color="0000EE", underline="single")

    for col_idx in range(1, len(headers) + 1):
        max_len = max((len(str(ws.cell(row=r, column=col_idx).value or ""))
                       for r in range(1, row_idx + 1)), default=10)
        ws.column_dimensions[get_column_letter(col_idx)].width = min(max(max_len, 12), 55)


def _sheet_summary(wb, agents):
    ws = wb.create_sheet("Summary Statistics")

    def write_table(start_col, title, rows, header_color):
        ws.cell(row=1, column=start_col, value=title).font = _font(bold=True, size=11, color=header_color)
        ws.cell(row=1, column=start_col).fill = _fill("F5F5F5")
        ws.merge_cells(start_row=1, start_column=start_col,
                       end_row=1, end_column=start_col + len(rows[0]) - 1)
        for col_offset, header in enumerate(rows[0]):
            cell        = ws.cell(row=2, column=start_col + col_offset, value=header)
            cell.font   = _font(bold=True, size=10)
            cell.fill   = _fill("E8E8E8")
            cell.border = THIN_BORDER
        for row_offset, row_data in enumerate(rows[1:], 3):
            for col_offset, val in enumerate(row_data):
                cell        = ws.cell(row=row_offset, column=start_col + col_offset, value=val)
                cell.border = THIN_BORDER
                cell.font   = _font(size=10)
        ws.column_dimensions[get_column_letter(start_col)].width = 22

    total = max(len(agents), 1)
    beh   = len([a for a in agents if a.advantage == "behavioral"])
    ana   = len([a for a in agents if a.advantage == "analytical"])
    inf   = len([a for a in agents if a.advantage == "informational"])
    write_table(1, "By Advantage", [
        ["Advantage", "Count", "%"],
        ["Informational", inf, f"{inf/total*100:.0f}%"],
        ["Analytical",    ana, f"{ana/total*100:.0f}%"],
        ["Behavioral",    beh, f"{beh/total*100:.0f}%"],
        ["Total",         total, "100%"],
    ], header_color="EF9F27")

    cx_rows = [["Complexity", "Count", "%"]]
    for ck, cl in COMPLEXITIES:
        count = len([a for a in agents if a.complexity == ck])
        cx_rows.append([cl.split("(")[0].strip(), count, f"{count/total*100:.0f}%"])
    cx_rows.append(["Total", total, "100%"])
    write_table(5, "By Complexity", cx_rows, header_color="1D9E75")

    stage_rows = [["Stage", "Count"]]
    for sk, sl in STAGES:
        count = len([a for a in agents if sk in a.stages_list])
        stage_rows.append([sl, count])
    write_table(9, "By Stage", stage_rows, header_color="1E3A5F")

    low    = len([a for a in agents if a.autonomy == "low"])
    medium = len([a for a in agents if a.autonomy == "medium"])
    high   = len([a for a in agents if a.autonomy == "high"])
    write_table(13, "By Autonomy", [
        ["Autonomy", "Count", "%"],
        ["Low",    low,    f"{low/total*100:.0f}%"],
        ["Medium", medium, f"{medium/total*100:.0f}%"],
        ["High",   high,   f"{high/total*100:.0f}%"],
        ["Total",  total,  "100%"],
    ], header_color="D85A30")


# ─── Word XML helpers ─────────────────────────────────────────────────────────

def _set_paragraph_spacing(para, before=0, after=6):
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after  = Pt(after)


def _set_run_color(run, hex_color):
    r = int(hex_color[0:2], 16)
    g = int(hex_color[2:4], 16)
    b = int(hex_color[4:6], 16)
    run.font.color.rgb = RGBColor(r, g, b)


def _add_hyperlink(para, url, text):
    part  = para.part
    r_id  = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rs  = OxmlElement("w:rStyle")
    rs.set(qn("w:val"), "Hyperlink")
    rPr.append(rs)
    run.append(rPr)
    t      = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hl.append(run)
    para._p.append(hl)


def _add_internal_link(para, anchor, text, color="1155CC", size=11):
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("w:anchor"), anchor)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rf  = OxmlElement("w:rFonts")
    rf.set(qn("w:ascii"), "Calibri")
    rf.set(qn("w:hAnsi"), "Calibri")
    rPr.append(rf)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    rPr.append(sz)
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    rPr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    run.append(rPr)
    t = OxmlElement("w:t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    run.append(t)
    hl.append(run)
    para._p.append(hl)


def _add_bookmark(para, name, bm_id):
    bm_start = OxmlElement("w:bookmarkStart")
    bm_start.set(qn("w:id"), str(bm_id))
    bm_start.set(qn("w:name"), name)
    bm_end = OxmlElement("w:bookmarkEnd")
    bm_end.set(qn("w:id"), str(bm_id))
    para._p.insert(0, bm_start)
    para._p.append(bm_end)


def _add_pageref_field(para, bookmark_name, color="555555", size=10):
    """Insert a PAGEREF field that resolves to the bookmark's page when opened in Word."""
    def _rpr():
        rPr = OxmlElement("w:rPr")
        rf  = OxmlElement("w:rFonts")
        rf.set(qn("w:ascii"), "Calibri")
        rPr.append(rf)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(size * 2))
        rPr.append(sz)
        ce = OxmlElement("w:color")
        ce.set(qn("w:val"), color)
        rPr.append(ce)
        return rPr

    for fld_type, content in [
        ("begin",    None),
        ("instrText", f" PAGEREF {bookmark_name} \\h "),
        ("separate", None),
        ("cached",   "?"),
        ("end",      None),
    ]:
        r = OxmlElement("w:r")
        r.append(_rpr())
        if fld_type == "instrText":
            it = OxmlElement("w:instrText")
            it.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            it.text = content
            r.append(it)
        elif fld_type == "cached":
            t = OxmlElement("w:t")
            t.text = content
            r.append(t)
        else:
            fc = OxmlElement("w:fldChar")
            fc.set(qn("w:fldCharType"), fld_type if fld_type != "cached" else "end")
            if fld_type == "begin":
                fc.set(qn("w:dirty"), "1")
            r.append(fc)
        para._p.append(r)


def _set_cell_shading(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:shd")):
        tcPr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_cell_valign(cell, val="center"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:vAlign")):
        tcPr.remove(existing)
    va = OxmlElement("w:vAlign")
    va.set(qn("w:val"), val)
    tcPr.append(va)


def _set_cell_borders(cell, color_hex, size="4"):
    """Set all four borders of a cell to a specific color."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(existing)
    tcBorders = OxmlElement("w:tcBorders")
    for name in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color_hex)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _remove_table_borders(tbl):
    tblPr = tbl._tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl._tbl.insert(0, tblPr)
    for existing in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(existing)
    tblBorders = OxmlElement("w:tblBorders")
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:val"), "none")
        tblBorders.append(el)
    tblPr.append(tblBorders)

    for row in tbl.rows:
        for cell in row.cells:
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            for existing in tcPr.findall(qn("w:tcBorders")):
                tcPr.remove(existing)
            tcBorders = OxmlElement("w:tcBorders")
            for name in ("top", "left", "bottom", "right"):
                el = OxmlElement(f"w:{name}")
                el.set(qn("w:val"), "none")
                tcBorders.append(el)
            tcPr.append(tcBorders)


def _set_table_borders_color(tbl, color_hex="FFFFFF"):
    """Set all table and cell borders to the given color."""
    tblPr = tbl._tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl._tbl.insert(0, tblPr)
    for existing in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(existing)
    tblBorders = OxmlElement("w:tblBorders")
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color_hex)
        tblBorders.append(el)
    tblPr.append(tblBorders)

    for row in tbl.rows:
        for cell in row.cells:
            _set_cell_borders(cell, color_hex)


def _set_col_widths(tbl, widths_twips):
    tbl_el = tbl._tbl
    for existing in tbl_el.findall(qn("w:tblGrid")):
        tbl_el.remove(existing)
    tblGrid = OxmlElement("w:tblGrid")
    for w in widths_twips:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        tblGrid.append(gc)
    tblPr = tbl_el.find(qn("w:tblPr"))
    if tblPr is not None:
        tblPr.addnext(tblGrid)
    else:
        tbl_el.insert(0, tblGrid)
    for row in tbl.rows:
        for col_idx, cell in enumerate(row.cells):
            if col_idx < len(widths_twips):
                tc   = cell._tc
                tcPr = tc.get_or_add_tcPr()
                for existing in tcPr.findall(qn("w:tcW")):
                    tcPr.remove(existing)
                tcW = OxmlElement("w:tcW")
                tcW.set(qn("w:w"),    str(widths_twips[col_idx]))
                tcW.set(qn("w:type"), "dxa")
                tcPr.append(tcW)


def _set_row_height(row, height_twips, exact=True):
    tr   = row._tr
    trPr = tr.find(qn("w:trPr"))
    if trPr is None:
        trPr = OxmlElement("w:trPr")
        tr.insert(0, trPr)
    for existing in trPr.findall(qn("w:trHeight")):
        trPr.remove(existing)
    trH = OxmlElement("w:trHeight")
    trH.set(qn("w:val"),   str(height_twips))
    trH.set(qn("w:hRule"), "exact" if exact else "atLeast")
    trPr.append(trH)


def _word_nested_timeline(host_cell, agent_stages, adv_color):
    """Build a dot-timeline nested table inside a Word table cell."""
    active_set = set(agent_stages)

    # widths: 7 dot cols (302 twips each) + 6 connector cols (120 twips each) = 2834 twips ≈ 5 cm
    DOT_W, CONN_W = 302, 120
    widths = []
    for i in range(7):
        widths.append(DOT_W)
        if i < 6:
            widths.append(CONN_W)

    # Remove host cell's default paragraph so only the nested table remains
    for p in list(host_cell.paragraphs):
        p._element.getparent().remove(p._element)

    nested = host_cell.add_table(rows=2, cols=13)
    nested.style = "Table Grid"
    _remove_table_borders(nested)
    _set_col_widths(nested, widths)
    _set_row_height(nested.rows[0], 280, exact=False)   # label row — auto-expands
    _set_row_height(nested.rows[1], 320, exact=True)    # dot row

    for stage_idx, (stage_key, _) in enumerate(STAGES):
        col       = stage_idx * 2
        is_active = stage_key in active_set

        # Label row
        lp = nested.cell(0, col).paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_spacing(lp, before=0, after=0)
        if is_active:
            lr = lp.add_run(STAGE_ABBREV.get(stage_key, stage_key))
            lr.font.name = "Calibri"
            lr.font.size = Pt(6)
            lr.font.bold = True
            _set_run_color(lr, "9D7D6F")

        # Dot row
        dp = nested.cell(1, col).paragraphs[0]
        dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_spacing(dp, before=0, after=0)
        dr = dp.add_run("●")
        dr.font.size = Pt(9)
        _set_run_color(dr, "9D7D6F" if is_active else "CCCCCC")

        # Connector
        if stage_idx < 6:
            conn_col = col + 1
            _set_paragraph_spacing(nested.cell(0, conn_col).paragraphs[0], before=0, after=0)
            cp = nested.cell(1, conn_col).paragraphs[0]
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_paragraph_spacing(cp, before=0, after=0)
            cr = cp.add_run("─")
            cr.font.size = Pt(8)
            _set_run_color(cr, "CCCCCC")


def _word_agent_header(doc, num, agent, adv_color, cx_color):
    """
    Add a 4-column borderless header table that matches the screenshot:
      [N. Agent Title]  [COMPLEXITY BADGE]  [timeline]  [ADV BADGE]
    """
    adv_label_short = {"behavioral": "BEH", "analytical": "ANA",
                       "informational": "INF"}.get(agent.advantage, "—")
    cx_short = COMPLEXITY_SHORT.get(agent.complexity, agent.complexity or "—")

    # Column widths in twips (total ≈ 16 cm = 9072 twips)
    # [6.5cm | 2.3cm | 5.5cm | 1.7cm]
    col_widths = [3685, 1304, 3118, 964]

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    _remove_table_borders(tbl)
    _set_col_widths(tbl, col_widths)
    _set_row_height(tbl.rows[0], 800, exact=False)

    cells = tbl.rows[0].cells

    # ── Col 0: Agent title ────────────────────────────────────────────────
    _set_cell_valign(cells[0], "center")
    title_para = cells[0].paragraphs[0]
    _set_paragraph_spacing(title_para, before=0, after=0)
    title_run = title_para.add_run(f"{num}. {agent.name}")
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    _set_run_color(title_run, "1E3A5F")

    # ── Col 1: Complexity badge ───────────────────────────────────────────
    cx_bg      = COMPLEXITY_BADGE_BG.get(agent.complexity, "3D3D3D")
    is_white_swan = agent.complexity == "white"
    _set_cell_shading(cells[1], cx_bg)
    if is_white_swan:
        _set_cell_borders(cells[1], "000000", "6")
    _set_cell_valign(cells[1], "center")
    cx_para = cells[1].paragraphs[0]
    cx_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(cx_para, before=0, after=0)
    cx_run = cx_para.add_run(cx_short)
    cx_run.font.name   = "Calibri"
    cx_run.font.size   = Pt(8)
    cx_run.font.bold   = True
    cx_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00) if is_white_swan else RGBColor(0xFF, 0xFF, 0xFF)

    # ── Col 2: Timeline ───────────────────────────────────────────────────
    _set_cell_valign(cells[2], "center")
    _word_nested_timeline(cells[2], agent.stages_list, adv_color)

    # ── Col 3: Advantage badge ────────────────────────────────────────────
    _set_cell_shading(cells[3], adv_color)
    _set_cell_valign(cells[3], "center")
    adv_para = cells[3].paragraphs[0]
    adv_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(adv_para, before=0, after=0)
    adv_run = adv_para.add_run(adv_label_short)
    adv_run.font.name   = "Calibri"
    adv_run.font.size   = Pt(11)
    adv_run.font.bold   = True
    adv_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


# ─── Word export ─────────────────────────────────────────────────────────────

def _enable_update_fields(doc):
    """Tell Word to recalculate all fields (PAGEREF, etc.) when the document is opened."""
    settings_el = doc.settings._element
    for existing in settings_el.findall(qn("w:updateFields")):
        settings_el.remove(existing)
    uf = OxmlElement("w:updateFields")
    uf.set(qn("w:val"), "1")
    settings_el.append(uf)


def build_word(agents):
    doc = Document()

    for section in doc.sections:
        section.page_height   = Cm(29.7)
        section.page_width    = Cm(21.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    sorted_agents = _sort_by_stage(agents)
    agent_numbers = {a.id: i + 1 for i, a in enumerate(sorted_agents)}
    stage_map     = dict(STAGES)
    bm_counter    = [0]

    # ── Cover page ────────────────────────────────────────────────────────────

    # Title
    tp = doc.add_paragraph()
    tr = tp.add_run("AI Agents in Investment Management")
    tr.font.bold = True
    tr.font.size = Pt(20)
    tr.font.name = "Calibri"
    _set_run_color(tr, "1E3A5F")
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(tp, before=0, after=6)

    # Subtitle
    sp = doc.add_paragraph()
    sr = sp.add_run("Multi-Dimensional Classification Report using")
    sr.font.size   = Pt(13)
    sr.font.italic = True
    sr.font.name   = "Calibri"
    sp.alignment   = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(sp, before=0, after=0)

    sp2 = doc.add_paragraph()
    sr2 = sp2.add_run("Panthera's 3D classification framework")
    sr2.font.size      = Pt(13)
    sr2.font.italic    = True
    sr2.font.underline = True
    sr2.font.name      = "Calibri"
    _set_run_color(sr2, "1155CC")
    sp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(sp2, before=0, after=10)

    # Date
    dp = doc.add_paragraph()
    dr = dp.add_run(datetime.utcnow().strftime("%d %B %Y"))
    dr.font.size = Pt(10)
    dr.font.name = "Calibri"
    dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(dp, before=0, after=20)

    # "Agents list :" label
    al = doc.add_paragraph()
    alr = al.add_run("Agents list :")
    alr.font.name = "Calibri"
    alr.font.size = Pt(11)
    _set_paragraph_spacing(al, before=0, after=8)

    # Group by first pipeline stage
    grouped: dict = {}
    for a in sorted_agents:
        first = next((sk for sk in STAGE_KEYS if sk in a.stages_list), None)
        if first:
            grouped.setdefault(first, []).append(a)

    stages_present = [sk for sk in STAGE_KEYS if sk in grouped]
    mid            = (len(stages_present) + 1) // 2
    left_stages    = stages_present[:mid]
    right_stages   = stages_present[mid:]

    def _fill_toc_column(cell, stage_keys):
        """Populate a TOC column cell with stage headings and agent entries."""
        for s_idx, sk in enumerate(stage_keys):
            # Stage heading
            sp = cell.add_paragraph()
            sr = sp.add_run(f"{stage_map[sk]} :")
            sr.font.bold   = True
            sr.font.italic = True
            sr.font.name   = "Calibri"
            sr.font.size   = Pt(11)
            sp.paragraph_format.space_before = Pt(8 if s_idx > 0 else 0)
            sp.paragraph_format.space_after  = Pt(2)

            for agent in grouped[sk]:
                num = agent_numbers[agent.id]
                bm  = f"agent_{agent.id}"

                ep = cell.add_paragraph()
                ep.paragraph_format.left_indent  = Cm(0.8)
                ep.paragraph_format.space_before = Pt(2)
                ep.paragraph_format.space_after  = Pt(2)

                nr = ep.add_run(f"{num}. ")
                nr.font.name = "Calibri"
                nr.font.size = Pt(10.5)

                _add_internal_link(ep, bm, agent.name, size=10.5)

                dot_run = ep.add_run("  " + "─" * 30 + "  p.")
                dot_run.font.name      = "Calibri"
                dot_run.font.size      = Pt(7)
                dot_run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)

                _add_pageref_field(ep, bm, color="444444", size=10)

    # 2-column TOC table (no borders, full width)
    toc_tbl = doc.add_table(rows=1, cols=2)
    toc_tbl.style = "Table Grid"
    _remove_table_borders(toc_tbl)
    _set_col_widths(toc_tbl, [4536, 4536])   # 8cm each = 16cm total

    _fill_toc_column(toc_tbl.rows[0].cells[0], left_stages)
    _fill_toc_column(toc_tbl.rows[0].cells[1], right_stages)

    doc.add_page_break()

    # ── Agent sections ────────────────────────────────────────────────────────

    complexity_map = dict(zip(COMPLEXITY_KEYS, COMPLEXITY_LABELS))

    for agent_idx, agent in enumerate(sorted_agents):
        num       = agent_numbers[agent.id]
        bm        = f"agent_{agent.id}"
        adv_color = _advantage_solid(agent.advantage)
        cx_color  = COMPLEXITY_BADGE_BG.get(agent.complexity, "888888")

        # Invisible bookmark paragraph (for PAGEREF to resolve)
        bm_anchor = doc.add_paragraph()
        _add_bookmark(bm_anchor, bm, bm_counter[0])
        bm_counter[0] += 1
        bm_anchor.paragraph_format.space_before = Pt(0)
        bm_anchor.paragraph_format.space_after  = Pt(0)

        # ── 4-col header: title | cx_badge | timeline | adv_badge ────────────
        _word_agent_header(doc, num, agent, adv_color, cx_color)

        # URL below header (if exists)
        if agent.url:
            url_para = doc.add_paragraph()
            _add_hyperlink(url_para, agent.url, agent.url)
            url_para.paragraph_format.space_before = Pt(2)
            url_para.paragraph_format.space_after  = Pt(4)

        # Italic description
        if agent.description:
            desc_p = doc.add_paragraph()
            desc_r = desc_p.add_run(agent.description)
            desc_r.font.italic = True
            desc_r.font.name   = "Calibri"
            desc_r.font.size   = Pt(10)
            _set_paragraph_spacing(desc_p, before=4, after=8)

        # ── Classification table ──────────────────────────────────────────────
        adv_label    = {"behavioral": "Behavioral", "analytical": "Analytical",
                        "informational": "Informational"}.get(agent.advantage, "—")
        cx_label     = complexity_map.get(agent.complexity, "—")
        stages_label = ", ".join(stage_map.get(s, s) for s in agent.stages_list) or "—"
        auto_label   = (agent.autonomy or "—").title()
        type_label   = (agent.agent_type or "—").replace("-", " ").title()

        tbl = doc.add_table(rows=6, cols=3)
        tbl.style = "Table Grid"

        for cell, txt in zip(tbl.rows[0].cells, ["Dimension", "Classification", "Rationale"]):
            cell.text = txt
            p = cell.paragraphs[0]
            p.runs[0].font.bold = True
            p.runs[0].font.name = "Calibri"
            p.runs[0].font.size = Pt(9.5)
            _set_run_color(p.runs[0], "9D7D6F")
            _set_cell_shading(cell, "FAFAFA")

        rd = agent.rationale_dict

        # (dim_label, cls_text, color_square_hex, rationale_text)
        data_rows = [
            ("Investment Process Stage(s)", stages_label, None,      rd.get("stages", "")),
            ("Complexity Range",            cx_label,     cx_color,  rd.get("complexity", "")),
            ("Comparative Advantage",       adv_label,    adv_color, rd.get("advantage", "")),
            ("Autonomy Level",              auto_label,   None,      rd.get("autonomy", "")),
            ("Agent Type",                  type_label,   None,      rd.get("agent_type", "")),
        ]

        for row_idx, (dim_txt, cls_txt, color_sq, rat_txt) in enumerate(data_rows, 1):
            cells = tbl.rows[row_idx].cells
            bg    = "F5F5F5" if row_idx % 2 == 0 else "FFFFFF"

            cells[0].text = dim_txt
            for run in cells[0].paragraphs[0].runs:
                run.font.name = "Calibri"
                run.font.size = Pt(9)

            cls_para = cells[1].paragraphs[0]
            if color_sq:
                sq_char = "□  " if color_sq == "FFFFFF" else "■  "
                sq_hex  = "000000" if color_sq == "FFFFFF" else color_sq
                sq = cls_para.add_run(sq_char)
                sq.font.name = "Calibri"
                sq.font.size = Pt(9)
                _set_run_color(sq, sq_hex)
            cr = cls_para.add_run(cls_txt)
            cr.font.name = "Calibri"
            cr.font.size = Pt(9)

            _set_cell_valign(cells[2], "top")
            rat_run = cells[2].paragraphs[0].add_run(rat_txt)
            rat_run.font.name   = "Calibri"
            rat_run.font.size   = Pt(8.5)
            rat_run.font.italic = True
            _set_run_color(rat_run, "555555")

            for cell in cells:
                _set_cell_shading(cell, bg)

        for row in tbl.rows:
            row.cells[0].width = Cm(3.5)
            row.cells[1].width = Cm(4.0)
            row.cells[2].width = Cm(8.5)

        _set_table_borders_color(tbl, "FFFFFF")

        _set_paragraph_spacing(doc.add_paragraph(), before=0, after=2)

        # KEY FEATURES
        h3_feat = doc.add_heading("KEY FEATURES", level=3)
        for run in h3_feat.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            run.font.bold = True
            _set_run_color(run, adv_color)
        _set_paragraph_spacing(h3_feat, before=6, after=3)

        for feature in agent.features_list:
            fp = doc.add_paragraph(feature, style="List Bullet")
            for run in fp.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(10)

        # Separator between agents
        if agent_idx < len(sorted_agents) - 1:
            hr_para = doc.add_paragraph()
            hr_para.paragraph_format.space_before = Pt(14)
            hr_para.paragraph_format.space_after  = Pt(14)
            pPr    = hr_para._p.get_or_add_pPr()
            pBdr   = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"),   "single")
            bottom.set(qn("w:sz"),    "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "CCCCCC")
            pBdr.append(bottom)
            pPr.append(pBdr)

    _enable_update_fields(doc)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ─── PDF – ReportLab helpers ─────────────────────────────────────────────────

def _hex(hex_str):
    return HexColor(f"#{hex_str}")


class _TimelineFlowable(Flowable):
    """Minimalist dot-line timeline for N stages."""

    def __init__(self, stages, active_stages, adv_hex, width=5.5 * cm, height=1.1 * cm):
        super().__init__()
        self._stages        = stages
        self._active        = set(active_stages)
        self._adv_color     = _hex(adv_hex)
        self.width          = width
        self.height         = height

    def wrap(self, aw, ah):
        return self.width, self.height

    def draw(self):
        c      = self.canv
        n      = len(self._stages)
        margin = 12
        span   = self.width - 2 * margin
        step   = span / (n - 1) if n > 1 else span
        y_dot  = 10
        y_lbl  = y_dot + 8
        r_dot  = 4

        # Connecting line
        c.setStrokeColor(_hex("CCCCCC"))
        c.setLineWidth(1.5)
        c.line(margin, y_dot, self.width - margin, y_dot)

        for i, (sk, _) in enumerate(self._stages):
            x         = margin + i * step
            is_active = sk in self._active
            col       = self._adv_color if is_active else _hex("CCCCCC")

            # Dot
            c.setFillColor(col)
            c.setStrokeColor(col)
            c.circle(x, y_dot, r_dot, fill=1, stroke=0)

            # Label above active dots
            if is_active:
                abbrev = STAGE_ABBREV.get(sk, sk)
                c.setFont("Helvetica-Bold", 5.5)
                c.setFillColor(self._adv_color)
                c.drawCentredString(x, y_lbl, abbrev)


def _pdf_styles(adv_hex):
    """Return a dict of ParagraphStyles keyed by purpose."""
    adv_color  = _hex(adv_hex)
    base       = getSampleStyleSheet()["Normal"]
    base.fontName = "Helvetica"

    def make(name, **kw):
        return ParagraphStyle(name, parent=base, **kw)

    return {
        "title":       make("title",       fontName="Helvetica-Bold",  fontSize=20,
                            textColor=_hex("1E3A5F"), alignment=TA_CENTER,
                            leading=26, spaceAfter=10),
        "subtitle1":   make("subtitle1",   fontName="Helvetica-Oblique", fontSize=13,
                            textColor=_hex("333333"), alignment=TA_CENTER,
                            leading=18, spaceAfter=2),
        "subtitle2":   make("subtitle2",   fontName="Helvetica-Oblique", fontSize=13,
                            textColor=_hex("1155CC"), alignment=TA_CENTER,
                            underline=1, spaceAfter=8),
        "date":        make("date",        fontSize=10, alignment=TA_CENTER, spaceAfter=16),
        "toc_stage":   make("toc_stage",   fontName="Helvetica-Bold",  fontSize=11,
                            textColor=_hex("333333"), spaceAfter=2, spaceBefore=8),
        "toc_entry":   make("toc_entry",   fontSize=10.5, leftIndent=12,
                            spaceAfter=2, spaceBefore=2),
        "agent_label": make("agent_label", fontName="Helvetica-Bold",  fontSize=16,
                            textColor=_hex("1E3A5F"), leading=18),
        "cx_badge":    make("cx_badge",    fontName="Helvetica-Bold",  fontSize=8,
                            textColor=white, alignment=TA_CENTER, leading=10),
        "adv_badge":   make("adv_badge",   fontName="Helvetica-Bold",  fontSize=14,
                            textColor=white, alignment=TA_CENTER),
        "url":         make("url",         fontSize=9,  textColor=_hex("1155CC"), spaceAfter=4),
        "desc":        make("desc",        fontName="Helvetica-Oblique", fontSize=10,
                            spaceAfter=8, leading=14),
        "tbl_hdr":     make("tbl_hdr",     fontName="Helvetica-Bold",  fontSize=9.5,
                            textColor=_hex("9D7D6F")),
        "tbl_dim":     make("tbl_dim",     fontSize=9),
        "tbl_cls":     make("tbl_cls",     fontSize=9),
        "tbl_note":    make("tbl_note",    fontName="Helvetica-Oblique", fontSize=8.5,
                            textColor=_hex("555555")),
        "feat_head":   make("feat_head",   fontName="Helvetica-Bold",  fontSize=10,
                            textColor=adv_color, spaceBefore=8, spaceAfter=3),
        "feat_item":   make("feat_item",   fontSize=10, leftIndent=12,
                            bulletIndent=0, spaceAfter=2),
        "rat_head":    make("rat_head",    fontName="Helvetica-Bold",  fontSize=10,
                            textColor=_hex("555555"), spaceBefore=8, spaceAfter=3),
        "rat_body":    make("rat_body",    fontSize=10, spaceAfter=4),
    }


def _pdf_classification_table(agent, styles, cx_color, adv_color):
    stage_map      = dict(STAGES)
    complexity_map = dict(zip(COMPLEXITY_KEYS, COMPLEXITY_LABELS))

    adv_label    = {"behavioral": "Behavioral", "analytical": "Analytical",
                    "informational": "Informational"}.get(agent.advantage, "—")
    cx_label     = complexity_map.get(agent.complexity, "—")
    stages_label = ", ".join(stage_map.get(s, s) for s in agent.stages_list) or "—"
    auto_label   = (agent.autonomy or "—").title()
    type_label   = (agent.agent_type or "—").replace("-", " ").title()

    def hdr(txt):
        return Paragraph(txt, styles["tbl_hdr"])

    def dim(txt):
        return Paragraph(txt, styles["tbl_dim"])

    def cls_cell(txt, color_hex=None):
        if color_hex:
            return Paragraph(
                f'<font color="#{color_hex}">■</font>  {txt}',
                styles["tbl_cls"],
            )
        return Paragraph(txt, styles["tbl_cls"])

    rd = agent.rationale_dict

    rat_style = ParagraphStyle(
        "rat_inline", parent=styles["tbl_note"],
        fontName="Helvetica-Oblique", fontSize=8.5,
        textColor=_hex("555555"), leading=12,
    )

    def rat(key):
        return Paragraph(rd.get(key, ""), rat_style)

    data = [
        [hdr("Dimension"),                        hdr("Classification"),          hdr("Rationale")],
        [dim("Investment Process Stage(s)"),       cls_cell(stages_label),         rat("stages")],
        [dim("Complexity Range"),                  cls_cell(cx_label, cx_color),   rat("complexity")],
        [dim("Comparative Advantage"),             cls_cell(adv_label, adv_color), rat("advantage")],
        [dim("Autonomy Level"),                    cls_cell(auto_label),           rat("autonomy")],
        [dim("Agent Type"),                        cls_cell(type_label),           rat("agent_type")],
    ]

    col_w = [3.5 * cm, 4.0 * cm, 8.5 * cm]

    ts = TableStyle([
        ("GRID",          (0, 0), (-1, -1), 0.5, _hex("DDDDDD")),
        ("BACKGROUND",    (0, 0), (-1, 0),  _hex("FAFAFA")),
        ("BACKGROUND",    (0, 2), (1, 2),   _hex("F5F5F5")),
        ("BACKGROUND",    (0, 4), (1, 4),   _hex("F5F5F5")),
        ("VALIGN",        (0, 0), (-1, -1), "TOP"),
        ("VALIGN",        (0, 0), (-1, 0),  "MIDDLE"),
        ("TOPPADDING",    (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("LEFTPADDING",   (0, 0), (-1, -1), 6),
    ])

    return Table(data, colWidths=col_w, style=ts, repeatRows=1)


# ─── PDF export ───────────────────────────────────────────────────────────────

def build_pdf(agents):
    buf  = io.BytesIO()
    PAGE_W, PAGE_H = A4
    MARGIN = 2.5 * cm

    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=MARGIN, rightMargin=MARGIN,
        topMargin=MARGIN,  bottomMargin=MARGIN,
        title="AI Agents in Investment Management",
    )

    sorted_agents = _sort_by_stage(agents)
    agent_numbers = {a.id: i + 1 for i, a in enumerate(sorted_agents)}
    stage_map     = dict(STAGES)
    story         = []

    # ── Global styles (neutral – per-agent styles re-built with adv color) ────
    gst = _pdf_styles("888888")

    # ── Cover page ────────────────────────────────────────────────────────────

    story.append(Paragraph("AI Agents in Investment Management", gst["title"]))
    story.append(Paragraph("Multi-Dimensional Classification Report using", gst["subtitle1"]))
    story.append(Paragraph("<u>Panthera's 3D classification framework</u>", gst["subtitle2"]))
    story.append(Paragraph(datetime.utcnow().strftime("%d %B %Y"), gst["date"]))
    story.append(Spacer(1, 0.4 * cm))
    story.append(Paragraph("Agents list :", ParagraphStyle(
        "al", parent=gst["date"], alignment=TA_LEFT, fontName="Helvetica",
        fontSize=11, textColor=black, spaceAfter=6,
    )))

    # Group by first pipeline stage
    grouped: dict = {}
    for a in sorted_agents:
        first = next((sk for sk in STAGE_KEYS if sk in a.stages_list), None)
        if first:
            grouped.setdefault(first, []).append(a)

    stages_present = [sk for sk in STAGE_KEYS if sk in grouped]
    mid            = (len(stages_present) + 1) // 2
    left_stages    = stages_present[:mid]
    right_stages   = stages_present[mid:]

    def _toc_column_content(stage_keys):
        items = []
        for sk in stage_keys:
            items.append(Paragraph(f"{stage_map[sk]} :", ParagraphStyle(
                "ts", parent=gst["toc_stage"],
                spaceBefore=8 if items else 0,
            )))
            for a in grouped[sk]:
                num = agent_numbers[a.id]
                items.append(Paragraph(
                    f'  {num}. <a href="#{a.id}" color="#1155CC"><u>{a.name}</u></a>',
                    gst["toc_entry"],
                ))
        return items

    left_items  = _toc_column_content(left_stages)
    right_items = _toc_column_content(right_stages)
    max_rows    = max(len(left_items), len(right_items))
    left_items  += [Paragraph("", gst["toc_entry"])] * (max_rows - len(left_items))
    right_items += [Paragraph("", gst["toc_entry"])] * (max_rows - len(right_items))

    toc_data  = [[l, r] for l, r in zip(left_items, right_items)]
    usable_w  = PAGE_W - 2 * MARGIN
    toc_table = Table(toc_data, colWidths=[usable_w / 2, usable_w / 2])
    toc_table.setStyle(TableStyle([
        ("VALIGN",      (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING",(0, 0), (-1, -1), 4),
        ("TOPPADDING",  (0, 0), (-1, -1), 1),
        ("BOTTOMPADDING",(0, 0), (-1, -1), 1),
    ]))
    story.append(toc_table)
    story.append(PageBreak())

    # ── Agent sections ────────────────────────────────────────────────────────

    for agent_idx, agent in enumerate(sorted_agents):
        num       = agent_numbers[agent.id]
        adv_color = _advantage_solid(agent.advantage)
        cx_color  = COMPLEXITY_COLORS.get(agent.complexity, "888888")
        cx_short  = COMPLEXITY_SHORT.get(agent.complexity, "")
        adv_short = {"behavioral": "BEH", "analytical": "ANA",
                     "informational": "INF"}.get(agent.advantage, "—")
        st        = _pdf_styles(adv_color)

        # Invisible anchor for TOC links
        story.append(Paragraph(f'<a name="{agent.id}"/>', ParagraphStyle(
            "anchor", parent=st["agent_label"], fontSize=0, leading=0,
        )))

        # ── 4-column header table ─────────────────────────────────────────────
        title_para    = Paragraph(f"{num}. {agent.name}", st["agent_label"])
        cx_badge_para = Paragraph(cx_short, st["cx_badge"])
        adv_badge_para= Paragraph(adv_short, st["adv_badge"])
        timeline      = _TimelineFlowable(
            STAGES, agent.stages_list, "9D7D6F",
            width=5.5 * cm, height=1.1 * cm,
        )

        hdr_data  = [[title_para, cx_badge_para, timeline, adv_badge_para]]
        hdr_table = Table(
            hdr_data,
            colWidths=[6.5 * cm, 2.3 * cm, 5.5 * cm, 1.7 * cm],
            rowHeights=[1.4 * cm],
        )
        hdr_table.setStyle(TableStyle([
            ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
            ("BACKGROUND",    (1, 0), (1, 0),   _hex("3D3D3D")),
            ("BACKGROUND",    (3, 0), (3, 0),   _hex(adv_color)),
            ("ROUNDEDCORNERS",(1, 0), (1, 0),   [3, 3, 3, 3]),
            ("ROUNDEDCORNERS",(3, 0), (3, 0),   [3, 3, 3, 3]),
            ("LEFTPADDING",   (0, 0), (-1, -1), 6),
            ("RIGHTPADDING",  (0, 0), (-1, -1), 6),
            ("TOPPADDING",    (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        story.append(hdr_table)
        story.append(Spacer(1, 0.2 * cm))

        # URL
        if agent.url:
            story.append(Paragraph(
                f'<a href="{agent.url}" color="#1155CC"><u>{agent.url}</u></a>',
                st["url"],
            ))

        # Description
        if agent.description:
            story.append(Paragraph(agent.description, st["desc"]))

        # Classification table
        story.append(_pdf_classification_table(agent, st, cx_color, adv_color))
        story.append(Spacer(1, 0.25 * cm))

        # KEY FEATURES
        if agent.features_list:
            story.append(Paragraph("KEY FEATURES", st["feat_head"]))
            for feat in agent.features_list:
                story.append(Paragraph(f"• {feat}", st["feat_item"]))

        if agent_idx < len(sorted_agents) - 1:
            story.append(HRFlowable(
                width="100%", thickness=0.5,
                color=_hex("CCCCCC"),
                spaceBefore=14, spaceAfter=14,
            ))

    doc.build(story)
    buf.seek(0)
    return buf
